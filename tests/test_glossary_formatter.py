import pytest
from docx import Document
from glossary_formatter import format_glossary

# Define the test data
TEST_FILE_PATH = "test.docx"
TEST_PHRASE = "test phrase"
TEST_FONT_STYLE = "Arial"
TEST_BOLD = True
TEST_ITALIC = False


# Define the test fixtures
@pytest.fixture
def document_with_test_phrase():
    # Create a new Word document with the test phrase
    document = Document()
    document.add_paragraph(TEST_PHRASE)
    document.save(TEST_FILE_PATH)
    return document


# Define the parameterized test function
@pytest.mark.parametrize(
    "bold, italic, underline",
    [
        (True, False, False),
        (False, True, False),
        (False, False, True),
        (True, True, False),
        (True, False, True),
        (False, True, True),
        (True, True, True),
    ],
)
def test_format_glossary(document_with_test_phrase, bold, italic, underline):
    # Call the function to find and style the test phrase
    format_glossary(
        TEST_FILE_PATH,
        TEST_PHRASE,
        font_style=TEST_FONT_STYLE,
        make_bold=bold,
        make_italic=italic,
        underline=underline,
    )

    # Verify that the changes were made to the document
    document = Document(TEST_FILE_PATH)
    assert document.paragraphs[0].text == TEST_PHRASE
    assert document.paragraphs[0].runs[0].font.name == TEST_FONT_STYLE
    assert document.paragraphs[0].runs[0].bold == bold
    assert document.paragraphs[0].runs[0].italic == italic
    assert document.paragraphs[0].runs[0].underline == underline


# Define the test functions
def test_format_glossary_phrase_not_found():
    # Create a new Word document without the test phrase
    document = Document()
    document.add_paragraph("not the test phrase")
    document.save(TEST_FILE_PATH)

    # Call the function to find and style the test phrase
    format_glossary(
        TEST_FILE_PATH,
        TEST_PHRASE,
        font_style=TEST_FONT_STYLE,
        make_bold=TEST_BOLD,
        make_italic=TEST_ITALIC,
    )

    # Verify that the changes were not made to the document
    document = Document(TEST_FILE_PATH)
    assert document.paragraphs[0].text != TEST_PHRASE


def test_format_glossary_invalid_font_style():
    # Create a new Word document with the test phrase
    document = Document()
    document.add_paragraph(TEST_PHRASE)
    document.save(TEST_FILE_PATH)

    # Call the function with an invalid font style
    with pytest.raises(Exception):
        format_glossary(TEST_FILE_PATH, TEST_PHRASE, font_style="invalid font style")


def test_format_glossary_invalid_file_path():
    # Call the function with an invalid file path
    with pytest.raises(Exception):
        format_glossary(
            "invalid file path",
            TEST_PHRASE,
            font_style=TEST_FONT_STYLE,
            make_bold=TEST_BOLD,
            make_italic=TEST_ITALIC,
        )


def test_format_glossary_font_style_conflict():
    # Create a new Word document with the test phrase and an existing font style
    document = Document()
    paragraph = document.add_paragraph(TEST_PHRASE)
    paragraph.runs[0].font.name = "Times New Roman"
    document.save(TEST_FILE_PATH)

    # Call the function with a conflicting font style
    format_glossary(TEST_FILE_PATH, TEST_PHRASE, font_style=TEST_FONT_STYLE)

    # Verify that the specified font style was applied and did not override the existing style
    document = Document(TEST_FILE_PATH)
    assert document.paragraphs[0].runs[0].font.name == TEST_FONT_STYLE


def test_format_glossary_special_characters():
    # Create a new Word document with the test phrase and special characters
    document = Document()
    document.add_paragraph(
        "this is a " + TEST_PHRASE + " with special characters: !@#$%^&*()_+"
    )
    document.save(TEST_FILE_PATH)

    # Call the function to find and style the test phrase
    format_glossary(TEST_FILE_PATH, TEST_PHRASE, font_style=TEST_FONT_STYLE)

    # Verify that the changes were made to the document
    document = Document(TEST_FILE_PATH)
    assert document.paragraphs[0].runs[1].font.name == TEST_FONT_STYLE


def test_format_glossary_substring():
    # Create a new Word document with a longer phrase that contains the test phrase as a substring
    document = Document()
    document.add_paragraph(
        "this is a longer phrase that contains the " + TEST_PHRASE + " as a substring"
    )
    document.save(TEST_FILE_PATH)

    # Call the function to find and style the test phrase
    format_glossary(TEST_FILE_PATH, TEST_PHRASE, font_style=TEST_FONT_STYLE)

    # Verify that only the test phrase was found and styled
    document = Document(TEST_FILE_PATH)
    assert document.paragraphs[0].runs[1].font.name == TEST_FONT_STYLE
